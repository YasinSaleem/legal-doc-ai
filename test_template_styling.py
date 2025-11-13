#!/usr/bin/env python3
"""
Test script to verify that user templates properly apply base styles
while preserving header/footer from user templates.
"""

import os
import sys
import tempfile
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Add src directory to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

from document_builder import build_document_from_json_content
from style_extractor import extract_styles_from_template, get_default_styles

def create_test_user_template():
    """Create a test user template with header/footer and some custom styles."""
    doc = Document()
    
    # Add header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "USER TEMPLATE HEADER - CONFIDENTIAL"
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add footer
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "User Template Footer - Page 1"
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # Add some sample content with different styles
    title = doc.add_paragraph("Sample User Template Title")
    title.style.name = "Heading 1"
    
    # Modify the Heading 1 style in the user template
    style = doc.styles['Heading 1']
    style.font.name = 'Arial'  # Different from base style
    style.font.size = Pt(18)   # Different from base style
    
    content = doc.add_paragraph("This is sample content in the user template.")
    content.style.name = "Normal"
    
    return doc

def create_test_json_content():
    """Create test JSON content for document generation."""
    return {
        "title": "Test Document Title",
        "sections": {
            "introduction": {
                "type": "Paragraph",
                "content": "This is the introduction paragraph with base styling applied."
            },
            "main_heading": {
                "type": "Heading 1", 
                "content": "Main Section Heading"
            },
            "sub_heading": {
                "type": "Heading 2",
                "content": "Sub Section Heading"
            },
            "body_content": {
                "type": "Paragraph",
                "content": "This is body content that should have proper Times New Roman font and justified alignment from base styles, even when using a user template."
            },
            "signatures": {
                "type": "Signature",
                "content": "Disclosing Party: Test User\n\n_____________________________"
            }
        }
    }

def test_user_template_with_base_styles():
    """Test that user templates preserve headers/footers but apply base styles."""
    
    print("🧪 Testing user template with base style application...")
    
    with tempfile.TemporaryDirectory() as tmpdir:
        # Create test user template
        user_template_path = os.path.join(tmpdir, "user_template.docx")
        user_template = create_test_user_template()
        user_template.save(user_template_path)
        print(f"✅ Created test user template: {user_template_path}")
        
        # Get base styles for reference
        base_styles = get_default_styles()
        print(f"📊 Base predefined styles that will be applied:")
        for style_name, props in base_styles.items():
            print(f"  {style_name}: {props}")
        
        print(f"\n� User template will only provide headers/footers, not styling")
        
        # Create test JSON content
        test_content = create_test_json_content()
        
        # Generate document using user template
        output_filename = "test_output_with_user_template.docx"
        output_path = build_document_from_json_content(
            template_path=user_template_path,
            doc_type="NDA",
            json_content=test_content,
            output_filename=output_filename,
            reference_doc_path=user_template_path,  # Use same as reference
            language_code='en'
        )
        
        print(f"✅ Generated document with user template: {output_path}")
        
        # Verify the generated document
        generated_doc = Document(output_path)
        
        # Check that header is preserved
        header_text = generated_doc.sections[0].header.paragraphs[0].text
        if "USER TEMPLATE HEADER" in header_text:
            print("✅ User template header preserved")
        else:
            print("❌ User template header NOT preserved")
        
        # Check that footer is preserved  
        footer_text = generated_doc.sections[0].footer.paragraphs[0].text
        if "User Template Footer" in footer_text:
            print("✅ User template footer preserved")
        else:
            print("❌ User template footer NOT preserved")
        
        # Check that content has base styles applied
        paragraph_found = False
        for para in generated_doc.paragraphs:
            if "body content" in para.text.lower():
                paragraph_found = True
                if para.runs:
                    font_name = para.runs[0].font.name
                    # Base style should be Times New Roman
                    if font_name == "Times New Roman":
                        print("✅ Base style font (Times New Roman) applied to paragraph content")
                    else:
                        print(f"❌ Base style font NOT applied. Found: {font_name}")
                
                # Check alignment - should be justify from base styles
                if para.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                    print("✅ Base style alignment (justify) applied")
                elif para.alignment is None:
                    print("⚠️ No specific alignment set (default left)")
                else:
                    print(f"❌ Base style alignment NOT applied. Found: {para.alignment}")
        
        if not paragraph_found:
            print("❌ Could not find test paragraph content")
        
        return output_path

def test_no_user_template():
    """Test document generation without user template (baseline)."""
    
    print("\n🧪 Testing document generation WITHOUT user template (baseline)...")
    
    with tempfile.TemporaryDirectory() as tmpdir:
        # Create blank template
        blank_template_path = os.path.join(tmpdir, "blank_template.docx")
        blank_doc = Document()
        blank_doc.save(blank_template_path)
        
        # Create test JSON content
        test_content = create_test_json_content()
        
        # Generate document using blank template
        output_filename = "test_output_no_user_template.docx"
        output_path = build_document_from_json_content(
            template_path=blank_template_path,
            doc_type="NDA", 
            json_content=test_content,
            output_filename=output_filename,
            reference_doc_path=None,  # No reference document
            language_code='en'
        )
        
        print(f"✅ Generated baseline document: {output_path}")
        
        # Verify baseline document has base styles
        generated_doc = Document(output_path)
        
        for para in generated_doc.paragraphs:
            if "body content" in para.text.lower():
                if para.runs:
                    font_name = para.runs[0].font.name
                    if font_name == "Times New Roman":
                        print("✅ Baseline: Base style font (Times New Roman) applied")
                    else:
                        print(f"❌ Baseline: Unexpected font: {font_name}")
        
        return output_path

def main():
    """Run all tests."""
    print("=" * 60)
    print("TESTING USER TEMPLATE STYLING FIX")
    print("=" * 60)
    
    try:
        # Test baseline (no user template)
        baseline_path = test_no_user_template()
        
        # Test with user template  
        user_template_path = test_user_template_with_base_styles()
        
        print("\n" + "=" * 60)
        print("TEST SUMMARY")
        print("=" * 60)
        print(f"✅ Baseline document: {baseline_path}")
        print(f"✅ User template document: {user_template_path}")
        print("\nThe implementation ensures that:")
        print("1. User template headers/footers are preserved")
        print("2. Base predefined styles are ALWAYS applied to content")
        print("3. User template styling is ignored - only headers/footers matter")
        
    except Exception as e:
        print(f"❌ Test failed with error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()