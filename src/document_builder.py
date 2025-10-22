"""
document_builder.py
-------------------
Builds a Word document dynamically from:
- A prepared template (user-uploaded or base)
- Schema-defined structure (sections, headings, paragraphs)
- Styles extracted from style_extractor.py
- Placeholder content from content_generator.py

Header/footer elements are preserved.
"""

import os
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from style_extractor import load_style_json
from config import WORKING_DIR, SCHEMA_DIR


# ──────────────────────────────────────────────
# Helper: Map alignment string to WD_PARAGRAPH_ALIGNMENT
# ──────────────────────────────────────────────
ALIGN_MAP = {
    "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
    "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
    "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
    "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY
}


def apply_style_to_paragraph(paragraph, style_props: dict):
    """
    Apply style properties to a given paragraph.
    """
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.name = style_props.get("font", "Times New Roman")
    run.font.size = Pt(style_props.get("size", 12))
    run.font.bold = style_props.get("bold", False)
    run.font.italic = style_props.get("italic", False)
    run.font.underline = style_props.get("underline", False)

    alignment = ALIGN_MAP.get(style_props.get("align", "left"), WD_PARAGRAPH_ALIGNMENT.LEFT)
    paragraph.alignment = alignment

    # Spacing
    spacing = style_props.get("spacing", 1.0)
    paragraph.paragraph_format.line_spacing = spacing

    # Indentation
    left_indent = style_props.get("indent_left", 0)
    right_indent = style_props.get("indent_right", 0)
    paragraph.paragraph_format.left_indent = Pt(left_indent)
    paragraph.paragraph_format.right_indent = Pt(right_indent)


def add_signature_section(doc, signature_content: str, style_json: dict, doc_type: str = "NDA"):
    """Add a professional signature section based on document type."""
    from docx.shared import Inches
    
    # Parse signature content to extract party names
    lines = signature_content.split('\n')
    disclosing_party = ""
    receiving_party = ""
    
    for line in lines:
        if "Disclosing Party:" in line:
            disclosing_party = line.replace("Disclosing Party:", "").strip()
        elif "Receiving Party:" in line:
            receiving_party = line.replace("Receiving Party:", "").strip()
    
    signature_style = style_json.get("Signature", style_json.get("Normal", {}))
    
    if doc_type.upper() == "NDA":
        # NDA: Single signature for Disclosing Party only
        add_nda_signature(doc, disclosing_party, signature_style)
    else:
        # Contract and other documents: Side-by-side signatures for both parties
        add_contract_signature(doc, disclosing_party, receiving_party, signature_style)


def add_nda_signature(doc, disclosing_party: str, signature_style: dict):
    """Add NDA signature section (single party only)."""
    # Add disclosing party signature
    disclosing_para = doc.add_paragraph(f"Disclosing Party:\n{disclosing_party}\n\n_____________________________")
    apply_style_to_paragraph(disclosing_para, signature_style)
    
    # Add date below
    date_para = doc.add_paragraph("Date: _________________")
    apply_style_to_paragraph(date_para, signature_style)


def add_contract_signature(doc, disclosing_party: str, receiving_party: str, signature_style: dict):
    """Add Contract signature section (both parties side-by-side with invisible borders)."""
    from docx.shared import Inches
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    # Create a table for side-by-side signatures
    table = doc.add_table(rows=2, cols=2)
    
    # Remove table borders (make invisible)
    # table.style = 'Table Normal' (to be deleted)
    
    # Remove borders from all cells
    for row in table.rows:
        for cell in row.cells:
            # Remove cell borders
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)
            tcPr.append(tcBorders)
            
            # Apply styling to paragraphs in cell
            for paragraph in cell.paragraphs:
                apply_style_to_paragraph(paragraph, signature_style)
    
    # Set column widths
    table.columns[0].width = Inches(3)
    table.columns[1].width = Inches(3)
    
    # Add disclosing party signature
    disclosing_cell = table.cell(0, 0)
    disclosing_cell.text = f"Disclosing Party:\n{disclosing_party}\n\n_____________________________"
    
    # Add receiving party signature
    receiving_cell = table.cell(0, 1)
    receiving_cell.text = f"Receiving Party:\n{receiving_party}\n\n_____________________________"
    
    # Add date row spanning both columns
    date_cell = table.cell(1, 0)
    date_cell.merge(table.cell(1, 1))
    date_cell.text = "Date: _________________"


def build_document(template_path: str, doc_type: str, schema: list, output_filename: str) -> str:
    """
    Assemble a Word document based on schema and styles.
    Returns path to the saved draft.
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)

    # Remove existing body content, keep headers/footers
    doc._body.clear_content()

    # Load styles
    style_json = load_style_json(doc_type)

    # Iterate schema and add content
    for section in schema:
        sec_type = section.get("type", "Paragraph")
        sec_text = section.get("text", "")

        para = doc.add_paragraph(sec_text)
        style_props = style_json.get(sec_type, style_json.get("Paragraph", {}))
        apply_style_to_paragraph(para, style_props)

    # Ensure working directory exists
    os.makedirs(WORKING_DIR, exist_ok=True)
    output_path = os.path.join(WORKING_DIR, output_filename)
    doc.save(output_path)

    print(f"✅ Draft document created: {output_path}")
    return output_path

def build_document_from_json_content(template_path: str, doc_type: str, json_content: dict, output_filename: str, reference_doc_path: str = None) -> str:
    """
    Build a Word document from structured JSON content with enhanced styling.
    
    Args:
        template_path: Path to the base template
        doc_type: Type of document (NDA, Contract, etc.)
        json_content: Structured content with title and sections
        output_filename: Name for the output file
        reference_doc_path: Optional path to reference document for style extraction
    
    Returns:
        Path to the generated document
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)
    
    # Remove existing body content, keep headers/footers
    doc._body.clear_content()
    
    # Try to extract styles from reference document if provided
    style_json = None
    if reference_doc_path and os.path.exists(reference_doc_path):
        try:
            print(f"🎨 Attempting to extract styles from reference document: {reference_doc_path}")
            from style_extractor import extract_styles_from_template
            temp_style_path = os.path.join(WORKING_DIR, "temp_reference_styles.json")
            extract_styles_from_template(reference_doc_path, temp_style_path)
            
            with open(temp_style_path, "r", encoding="utf-8") as f:
                style_json = json.load(f)
            print("✅ Successfully extracted styles from reference document")
            
            # Clean up temp file
            os.remove(temp_style_path)
        except Exception as e:
            print(f"⚠️ Failed to extract styles from reference document: {e}")
            print("🔄 Falling back to predefined styles")
    
    # Fall back to predefined styles if extraction failed
    if not style_json:
        style_json = load_style_json(doc_type)
    
    # Add title if present (only once, not duplicated)
    title_added = False
    if "title" in json_content:
        title_para = doc.add_paragraph(json_content["title"])
        title_style = style_json.get("Heading 1", style_json.get("Normal", {}))
        apply_style_to_paragraph(title_para, title_style)
        title_added = True
    
    # Add sections
    if "sections" in json_content:
        for section_name, section_data in json_content["sections"].items():
            section_type = section_data.get("type", "Paragraph")
            section_content = section_data.get("content", "")
            
            # Skip duplicate title sections
            if title_added and section_type == "Heading 1" and section_content == json_content.get("title", ""):
                continue
            
            # Handle signature sections with special layout based on document type
            if section_type == "Signature":
                add_signature_section(doc, section_content, style_json, doc_type)
            else:
                # Add the section content
                para = doc.add_paragraph(section_content)
                
                # Apply appropriate styling with justified alignment for paragraphs
                style_props = style_json.get(section_type, style_json.get("Normal", {}))
                
                # Force justified alignment for all paragraph types
                if section_type in ["Paragraph", "Normal"]:
                    style_props = style_props.copy()  # Don't modify original
                    style_props["align"] = "justify"
                
                apply_style_to_paragraph(para, style_props)
    
    # Ensure output directory exists
    from config import DOC_OUTPUT_DIR
    os.makedirs(DOC_OUTPUT_DIR, exist_ok=True)
    output_path = os.path.join(DOC_OUTPUT_DIR, output_filename)
    doc.save(output_path)
    
    print(f"✅ Document built from JSON content: {output_path}")
    return output_path


# ──────────────────────────────────────────────
# Example manual test
# ──────────────────────────────────────────────
# if __name__ == "__main__":
#     # Example schema for NDA
#     test_schema = [
#         {"type": "Heading 1", "text": "Non-Disclosure Agreement"},
#         {"type": "Paragraph", "text": "[INTRODUCTION]"},
#         {"type": "Heading 2", "text": "Definitions"},
#         {"type": "Paragraph", "text": "[DEFINITIONS]"},
#         {"type": "Paragraph", "text": "[CONFIDENTIALITY]"},
#         {"type": "Paragraph", "text": "[TERM]"},
#         {"type": "Paragraph", "text": "[SIGNATURE]"}
#     ]

#     draft_path = build_document(
#         template_path="working/nda_working.docx",
#         doc_type="NDA",
#         schema=test_schema,
#         output_filename="nda_draft.docx"
#     )
