"""
placeholder_filler.py
---------------------
Replaces placeholders in a Word document with actual values
while preserving original formatting and styles.

Supports:
- Plain paragraphs
- Optional: tables (cells)
"""

import re
import os
from docx import Document
from config import OUTPUT_DIR

PLACEHOLDER_PATTERN = r"\[(.*?)\]"  # Matches [Name], [Company], etc.


def fill_placeholders(draft_path: str, data: dict, output_filename: str) -> str:
    """
    Replace placeholders in the draft document using the provided data dictionary.

    :param draft_path: Path to the draft Word document with placeholders
    :param data: Dictionary of replacements { "Name": "Alice", ... }
    :param output_filename: Filename for the final filled document
    :return: Path to the saved document
    """
    if not os.path.exists(draft_path):
        raise FileNotFoundError(f"Draft not found: {draft_path}")

    doc = Document(draft_path)

    # Replace placeholders in paragraphs
    for para in doc.paragraphs:
        inline_replace(para.runs, data)

    # Optional: Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    inline_replace(para.runs, data)

    # Ensure output folder exists
    output_dir = os.path.join(OUTPUT_DIR, "docs")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, output_filename)
    doc.save(output_path)

    print(f"✅ Placeholders filled. Final document saved: {output_path}")
    return output_path


def inline_replace(runs, data: dict):
    """
    Helper function to replace placeholders within runs of a paragraph.
    Preserves original formatting.
    """
    for run in runs:
        text = run.text
        matches = re.findall(PLACEHOLDER_PATTERN, text)
        for match in matches:
            key = match.strip()
            if key in data and data[key]:
                # Replace [KEY] with actual value
                text = text.replace(f"[{key}]", data[key])
        run.text = text


def detect_placeholders(doc_path: str) -> list:
    """
    Optional helper to list all placeholders in a document.
    """
    doc = Document(doc_path)
    placeholders = set()

    for para in doc.paragraphs:
        matches = re.findall(PLACEHOLDER_PATTERN, para.text)
        placeholders.update(matches)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = re.findall(PLACEHOLDER_PATTERN, cell.text)
                placeholders.update(matches)

    return list(placeholders)


# ──────────────────────────────────────────────
# Example manual test
# ──────────────────────────────────────────────
# if __name__ == "__main__":
#     test_data = {
#         "Name": "Alice Johnson",
#         "Company": "TechNova Ltd",
#         "Date": "October 15, 2025"
#     }
#     draft_file = "working/nda_draft.docx"
#     output_file = "Alice_Johnson_NDA.docx"

#     if os.path.exists(draft_file):
#         fill_placeholders(draft_file, test_data, output_file)
#     else:
#         print("❌ Draft file not found.")
